# %%
## いじるのはこのセルのパラメータのみでOK

file_name_download = "IPB-20210624.docx" #ダウンロードされる業績リストのファイル名
file_name_download_xlsx = "IPB-20210624.xlsx" #ダウンロードされるxlsxファイルのファイル名
file_name_download_check = "IPB-20210624check.docx" #集計チェック用のファイル名

globalmindate='2019-06-28' #これより後の業績を集める
globalmaxdate='2021-07-01' #これより前の業績を集める
smark='◎' #researchmapで課題番号紐づけありの論文にマーク付ける場合はここで指定。
ryoiki_linked = False #researchmapで課題番号紐づけありの論文のみ出力したい場合はTrue
allenglish = False #名前表記をすべて英語で統一する場合はTrue, 論文以外の名前表記を日本語にする場合False
SNfirst = False #英語名前表記をすべて名字先で統一する場合はTrue, 名字後で統一する場合はFalse
numberingPapers = True #出力の際に論文をナンバリング
peer_reviewed = True #査読ありのチェックが入った論文だけに限定する場合はTrue
firstnameInitial = True

sankodata=True # xlsxファイルの読み書きをする場合

docoutputpointsize=11 #11pt出力指定

# 領域メンバーの情報入りスプレッドシートのURL
sheeturl='https://docs.google.com/spreadsheets/d/1wce1XHSFGSBttupnSIqe_5abtijBb_hBYM2bfaV9Jn4/edit#gid=0' 

# 未入力のxlsxファイルのURL (githubにアップロード済)
blankxlsx='https://github.com/dbkk/docx-researchmap/blob/rev2021/inputfiles/R3%E4%B8%AD%E9%96%93%E8%A9%95%E4%BE%A1%E5%A0%B1%E5%91%8A%E6%9B%B8%EF%BC%881_%E9%A0%98%E5%9F%9F%E5%85%A8%E4%BD%93%EF%BC%89%EF%BC%8813%E5%8F%82%E8%80%83%E3%83%87%E3%83%BC%E3%82%BFExcel%E7%89%88%EF%BC%89.xlsx?raw=true'

# google formで取ったそれ以外のデータ (非公開) (readme.md参照)
sheeturlform='https://docs.google.com/spreadsheets/d/1WrQwU1QKW5KnJX8ORxt1dihAH3M_AAv15HfIVCW9TrA/edit?resourcekey#gid=1794423468'

#出力時の上限数
maxpap=99#9
maxtalk=99#5
maxsocial=99#5
maxmed=99#5
maxsonota=99#3
maxBSM=99#10

# %%
import requests,json,sys,os,gspread,time,re,openpyxl,datetime,xlrd
import numpy as np
import pandas as pd

if 'google.colab' in str(get_ipython()):
    %pip install python-docx
    from google.colab import files,auth
    from oauth2client.client import GoogleCredentials
    outputdirectory = ''
else:
    outputdirectory = '../docx-researchmap-outputs/' #ローカルで実行する場合は保存ファイルのディレクトリを適当に指定
    os.makedirs(outputdirectory,exist_ok=True)
from docx import Document
from docx.shared import Pt,Mm,RGBColor
from docx.enum.text import WD_UNDERLINE,WD_LINE_SPACING,WD_BREAK

file_name=outputdirectory+file_name_download
file_name_xlsx=outputdirectory+file_name_download_xlsx
file_name_check=outputdirectory+file_name_download_check

# %%
#スプレッドシートをダウンロード
sheeturl_csv=re.match("https://docs.google.com/spreadsheets/d/.+/",sheeturl).group(0)+"export?format=csv"
name_data=pd.read_csv(sheeturl_csv)
name_data

# %%
membernum=len(name_data)

if SNfirst:
    allnames=(name_data["Surname"]+' '+name_data["First name"]).to_list()
else:
    allnames=(name_data["First name"]+' '+name_data["Surname"]).to_list()
allSurname=name_data["Surname"].to_list()
allnamesJP=(name_data["苗字"]+" "+name_data["名"]).to_list()
allgroupnames=name_data["班"].to_list()
allgroupnum=name_data["番号"].to_list()
allmembers=name_data["researchmapID"].to_list()
allDB=name_data["代表分担協力"].values
allkeikaku=[b for a,b in zip(allgroupnames,allnamesJP) if a in ['A','B','C']]
allkeikakuPIs=[b for a,b,c in zip(allgroupnames,allnamesJP,allDB) if (a in ['A','B','C']) & (c =='D')]
allDaihyoBuntan=list(allDB)
allHan=(name_data["班"]+name_data["番号"].apply(str)).to_list()
grant_numbers=name_data["grantID"].to_list()
allmindate=name_data["Start date"].to_list()
allmaxdate=name_data["End date"].to_list()

#Exception names handling
altname2,altname3=name_data['著者名（2個目）'],name_data['著者名（3個目）']
arraltname2,arraltname3=altname2.values,altname3.values
nameList=allnames+list(arraltname2[~(pd.isna(altname2).values)])+list(arraltname3[~(pd.isna(altname3).values)])
nameList = [n.strip() for n in nameList]
daihyobuntanList=allDaihyoBuntan+list(allDB[~(pd.isna(altname2).values)])+list(allDB[~(pd.isna(altname3).values)])

# %%
# Function to set the name order
def SurnameFirst(namesDic,sn):
    oldnamelist=[]
    swap=0
    for indiv in namesDic:
        oldnamelist=oldnamelist+[indiv['name'].replace(',','').replace('.','')]
        #print(oldnamelist)
    return SurnameFirstList(oldnamelist,sn)

def SurnameFirstList(oldnamelist,sn):
    swap=0
    for name in oldnamelist:
        if sn in name.split(' '):
            if name.split(' ').index(sn)==0: # surname first
                swap= True ^ SNfirst
                break;
            else:
                swap= False ^ SNfirst
                break;
    if swap:
        newnamelist=[]
        for name in oldnamelist:
            namesplit=name.split(' ')
            names=[namesplit[-1]]+namesplit[:-1]
            newnamelist=newnamelist+[' '.join(names)]
    else:
        newnamelist=oldnamelist
    
    if SNfirst & firstnameInitial:
        holdlist=[]
        for name in newnamelist:
            namesplit=name.split(' ')
            names=[namesplit[0]]+[', ']+[namesplit[1][0]]+['.']
            holdlist=holdlist+[''.join(names)]
        newnamelist=holdlist
    elif firstnameInitial:
        holdlist=[]
        for name in newnamelist:
            namesplit=name.split(' ')
            sn=namesplit[-1]
            sn=sn.lower()
            sn=sn[0].upper()+sn[1:]
            names=[namesplit[0][0]]+['. ']+[sn]
            holdlist=holdlist+[''.join(names)]
        newnamelist=holdlist                    
    return newnamelist

def ReturnDictWOerror(dictdata,key,nodata):
    if key in dictdata.keys():
        return dictdata[key]
    else:
        return nodata

def ReturnDictContent(dictdata,key,key1,nodata=''):
    d=ReturnDictWOerror(dictdata,key,nodata)
    d1=ReturnDictWOerror(dictdata,key1,nodata)
    if d!=nodata:
        return d
    else:
        return d1

def commaR(vol,spage):
    if (vol=='') & (spage==''):
        return ''
    elif (vol=='') | (spage==''):
        return ' '
    else:
        return ', '

# %%
url = "https://api.researchmap.jp/"
itemslist = ["published_papers","research_projects","misc","presentations","books_etc","social_contribution","awards","media_coverage"]
jsonfiles={}
for name in allmembers:
    print('downloading: '+name)
    jsonfiles[name]={}
    for it in itemslist:
        r1 = requests.get(url+name+'/'+it)
        jsonfiles[name][it]=json.loads(r1.text)
        if 'error' in jsonfiles[name][it].keys():
            print(jsonfiles[name][it]['error'])
            print("  error in:"+it)

# %%
# make dictionary of all papers
i=0
PapersDict={}

doilist=[]
doiDict={}
titlelist=[]
titleDict={}
#dc = Document()
for ids,fullname,dh,mindate,maxdate,han in zip(allmembers,allnames,allDaihyoBuntan,allmindate,allmaxdate,allHan):
    surname=fullname.split(' ')[0 if SNfirst else 1]
    dfP = jsonfiles[ids]["published_papers"]
    dfG = jsonfiles[ids]["research_projects"]
    if 'items' in dfG.keys():
        grantID="0"
        for dfs in dfG['items']:
            if 'identifiers' in dfs.keys():
                if 'grant_number' in dfs['identifiers'].keys():
                    if dfs['identifiers']['grant_number'][0] in grant_numbers:
                        grantID=dfs['rm:id']
                        break
    if 'items' in dfP.keys():    
        for dfs in dfP['items']:
            if "authors" not in dfs.keys():
                continue
            if ('identifiers' in dfs.keys()) & (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
                doinum=[0]
                if 'doi' in dfs['identifiers'].keys():
                    doinum=dfs['identifiers']['doi']

                PapersDict[i]={}
                PapersDict[i]['issues']=False
                PapersDict[i]['preprint']=False
                correspo=False
                Ryoiki=False
                if 'rm:research_project_id' in dfs['identifiers'].keys():
                    if grantID in dfs['identifiers']['rm:research_project_id']:
                        Ryoiki=True
                        
                if "published_paper_owner_roles" in dfs.keys():
                    if ("corresponding" in dfs["published_paper_owner_roles"]) | ("last" in dfs["published_paper_owner_roles"]):
                        correspo=True

                jname=''        
                if "publication_name" in dfs.keys():
                    jname=ReturnDictContent(dfs["publication_name"],'en','ja','').upper()

                if jname =='ARXIV':
                    PapersDict[i]['preprint']=True
                    if "arxiv_id" in dfs['identifiers'].keys():
                        jname=dfs['identifiers']['arxiv_id'][0] + ' (preprint)'
                    else:
                        jname='arxiv'
                
                if not("publication_name" in dfs.keys()):
                    if "arxiv_id" in dfs['identifiers'].keys():
                        jname=dfs['identifiers']['arxiv_id'][0] + ' (preprint)'
                        PapersDict[i]['preprint']=True
                    elif doinum[0]!=0:
                        jname='DOI: '+doinum[0]
                        PapersDict[i]['preprint']=True
                    else:
                        jname='journal unspecified'
                        PapersDict[i]['issues']=True
                    
                Sname=SurnameFirst(ReturnDictContent(dfs["authors"],'en','ja',''),surname)

                spage=''
                if "starting_page" in dfs.keys():
                    if dfs["starting_page"]!='':
                        spage=dfs["starting_page"]

                vol=''
                if "volume" in dfs.keys():
                    if dfs["volume"]!='':
                        vol=' '+dfs["volume"]
                if doinum in doilist:
                    doiDict[doinum[0]]['name']=doiDict[doinum[0]]['name']+[fullname]
                    doiDict[doinum[0]]['Corresp']=doiDict[doinum[0]]['Corresp']+[correspo]
                else:
                    doiDict[doinum[0]]={}
                    doiDict[doinum[0]]['name']=[fullname]
                    doiDict[doinum[0]]['Corresp']=[correspo]
                    doiDict[doinum[0]]['count']=0
                    doilist=doilist+[doinum[0]]
                
                papertitle=ReturnDictContent(dfs['paper_title'],'en','ja','')
                papid=papertitle.upper().rstrip('.')

                if papid in titlelist:
                    titleDict[papid]['name'] = titleDict[papid]['name']+[fullname]
                    titleDict[papid]['Corresp'] = titleDict[papid]['Corresp']+[correspo]                    
                else:
                    titlelist = titlelist + [papid]
                    titleDict[papid] = {}
                    titleDict[papid]['name'] = [fullname]
                    titleDict[papid]['Corresp'] = [correspo]
                    titleDict[papid]['count']=0

                text1="\""+papertitle+"\"" +', '
                text2=jname+','+vol+commaR(vol,spage)+spage+ ' ('+dfs["publication_date"][:4] +').'
                if "description" in dfs.keys():
                    PapersDict[i]['oudan']=ReturnDictWOerror(dfs["description"],'ja','')
                else:
                    PapersDict[i]['oudan']=''
                #print(PapersDict[i]['oudan'])
                PapersDict[i]['kokunai']=False
                if "is_international_journal" in dfs.keys():
                    if not dfs["is_international_journal"]:
                        PapersDict[i]['kokunai']=True
                PapersDict[i]['text1']=text1
                PapersDict[i]['text2']=text2
                PapersDict[i]['papid']=papid
                PapersDict[i]['researcher']=fullname
                PapersDict[i]['authors']=Sname
                PapersDict[i]['date']=dfs["publication_date"]
                PapersDict[i]['referee']=ReturnDictContent(dfs,'referee','referee',False)
                PapersDict[i]['doi']=doinum[0]
                PapersDict[i]['ryoiki']=Ryoiki
                PapersDict[i]['Daihyo']=dh
                PapersDict[i]['han']=han
                PapersDict[i]['Corresp']=correspo
                i=i+1

# %%
# make dictionary of all talks, 'social_contribution', awards
TalksDict={}
SocialContDict={}
AwardsDict={}
i,j,l=0,0,0
for ids,fullname,fullnameJP,dh,mindate,maxdate,han in zip(allmembers,allnames,allnamesJP,allDaihyoBuntan,allmindate,allmaxdate,allHan):
    dfPr = jsonfiles[ids]["presentations"]
    dfSC = jsonfiles[ids]["social_contribution"]
    dfAw = jsonfiles[ids]["awards"]
    if 'items' in dfPr.keys():
        for dfs in dfPr['items']:
            if all([a in dfs.keys() for a in ["presentation_title","event",'publication_date','presenters']]):
                if (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
                    if ('en' in dfs["presenters"].keys()):
                        pname=dfs["presenters"]["en"][0]["name"]
                    else:
                        pname=dfs["presenters"]["ja"][0]["name"]
                    ename=ReturnDictContent(dfs["event"],'ja','en','')
                    ptitle=ReturnDictContent(dfs["presentation_title"],'ja','en','')
                    pdate=dfs["publication_date"]
                    TalksDict[i]={}
                    TalksDict[i]["presenter"]=fullnameJP
                    if allenglish:
                        TalksDict[i]['printname']=fullname
                    else:
                        TalksDict[i]["printname"]=fullnameJP
                    TalksDict[i]["event"]=ename
                    TalksDict[i]["presentation_title"]=ptitle
                    TalksDict[i]["date"]=pdate
                    TalksDict[i]["han"]=han
                    TalksDict[i]["invited"]=ReturnDictWOerror(dfs,'invited',False)
                    TalksDict[i]["international"]=ReturnDictWOerror(dfs,'is_international_presentation',False)
                    TalksDict[i]["keyoral"]= ReturnDictWOerror(dfs,"presentation_type",'')
                    i=i+1
    if 'items' in dfSC.keys():
        for dfs in dfSC['items']:
            if 'from_event_date' in dfs.keys():
                if (dfs['from_event_date']>=mindate) & (dfs['from_event_date']<=maxdate):
                    SocialContDict[j]={}
                    SocialContDict[j]['name']=fullnameJP
                    SocialContDict[j]["title"]=dfs['social_contribution_title']['ja']
                    SocialContDict[j]["date"]=dfs['from_event_date']
                    SocialContDict[j]["han"]=han
                    if 'event' in dfs.keys():
                        SocialContDict[j]["event"]=ReturnDictContent(dfs["event"],'ja','en','')
                    else:
                        SocialContDict[j]["event"]=''
                    j=j+1
    if 'items' in dfAw.keys():
        for dfs in dfAw['items']:
            if (dfs['award_date']>=mindate) & (dfs['award_date']<=maxdate):
                AwardsDict[l]={}
                AwardsDict[l]['name']=fullnameJP
                AwardsDict[l]['award_name']=ReturnDictContent(dfs['award_name'],'ja','en','')
                if 'association' in dfs.keys():
                    AwardsDict[l]['association']=ReturnDictContent(dfs['association'],'ja','en','')
                else:
                    AwardsDict[l]['association']=''
                AwardsDict[l]['award_date']=dfs['award_date']
                l=l+1

# %%
# make dictionary of all books_etc
booksDict={}
i=0
for ids,fullname,fullnameJP,dh,mindate,maxdate,han in zip(allmembers,allnames,allnamesJP,allDaihyoBuntan,allmindate,allmaxdate,allHan):
  dfM = jsonfiles[ids]["books_etc"]
  if 'items' in dfM.keys():
    for dfs in dfM['items']:
      if all([a in dfs.keys() for a in ['authors',"book_title","publication_date"]]):
        if (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
          if ('ja' in dfs["authors"].keys()):
              pname=dfs["authors"]["ja"][0]["name"]
          else:
              pname=dfs["authors"]["en"][0]["name"]
          ename=ReturnDictContent(dfs["book_title"],'ja','en','')
          if "book_owner_range" in dfs.keys():
            eoname=" \'"+ReturnDictContent(dfs["book_owner_range"],'ja','en','')+"\',"
          else:
            eoname=''
          if "book_owner_role" in dfs.keys():
            brole=" ("+dfs["book_owner_role"]+"),"
          else:
            brole=','
          if "publisher" in dfs.keys():
            pub=" "+ReturnDictContent(dfs["publisher"],'ja','en','')+","
          else:
            pub=''
          pdate=dfs["publication_date"]
          booksDict[i]={}
          booksDict[i]['authors']=fullname
          if allenglish:
            booksDict[i]['printname']=fullname
          else:
            booksDict[i]['printname']=fullnameJP
          booksDict[i]["book_title"]=' '+ename+','
          booksDict[i]["book_owner_role"]=brole
          booksDict[i]["book_owner_range"]=eoname
          booksDict[i]["publisher"]=pub
          booksDict[i]["date"]=pdate
          booksDict[i]["han"]=han
          i=i+1

# %%
# make dictionary of all MISCs
miscDict={}
medDict={}
i,j=0,0
for ids,fullname,fullnameJP,dh,mindate,maxdate,han in zip(allmembers,allnames,allnamesJP,allDaihyoBuntan,allmindate,allmaxdate,allHan):
  dfMis = jsonfiles[ids]["misc"]
  if 'items' in dfMis.keys():
    for dfs in dfMis['items']:
      if all([a in dfs.keys() for a in ['authors',"paper_title","publication_date","publication_name"]]):
        if  (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
          if ('ja' in dfs["authors"].keys()):
              pname=dfs["authors"]["ja"][0]["name"]
          else:
              pname=dfs["authors"]["en"][0]["name"]
          ename=ReturnDictContent(dfs["paper_title"],'ja','en','')
          ptitle=ReturnDictContent(dfs["publication_name"],'ja','en','')
          pdate=dfs["publication_date"]
          miscDict[i]={}
          miscDict[i]['authors']=fullname
          if allenglish:
            miscDict[i]['printname']=fullname
          else:
            miscDict[i]['printname']=fullnameJP

          miscDict[i]["paper_title"]=' \''+ename+'\','
          miscDict[i]["publication_name"]=' '+ptitle+','
          miscDict[i]["date"]=pdate
          miscDict[i]["han"]=han
          i=i+1
  dfMed = jsonfiles[ids]["media_coverage"]
  if 'items' in dfMed.keys():
    for dfs in dfMed['items']:
      if all([a in dfs.keys() for a in ["media_coverage_title","publication_date"]]):
        if  (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
          ename=ReturnDictContent(dfs["media_coverage_title"],'ja','en','')
          
          
          pdate=dfs["publication_date"]
          medDict[j]={}
          medDict[j]['authors']=fullname
          if allenglish:
            medDict[j]['printname']=fullname
          else:
            medDict[j]['printname']=fullnameJP
          medDict[j]["media_coverage_type"]=''
          if "media_coverage_type" in dfs.keys():
            medDict[j]["media_coverage_type"]=dfs["media_coverage_type"]
          ptitle=''
          if "publisher" in dfs.keys():
            ptitle=ReturnDictContent(dfs["publisher"],'ja','en','')+', '

          petitle=''
          if "event" in dfs.keys():
            petitle=ReturnDictContent(dfs["event"],'ja','en','')+','
#          print(' '+ptitle+petitle)

          pltitle=''
          if "location" in dfs.keys():
            pltitle=ReturnDictContent(dfs["location"],'ja','en','')+','
#          print(' '+ptitle+petitle+pltitle)

          medDict[j]["paper_title"]=' \''+ename+'\','
          medDict[j]["publication_name"]=' '+ptitle+petitle+pltitle
          medDict[j]["date"]=pdate
          medDict[j]["han"]=han
          j=j+1

# %%
# generate docx
document = Document()

section = document.sections[0]
section.left_margin = Mm(15)
section.right_margin = Mm(15)
section.top_margin = Mm(15)
section.bottom_margin = Mm(15)

nameListPrint=SurnameFirstList(nameList,'Okada')

if peer_reviewed:
    refbool=[True]
else:
    refbool=[True,False]

if ryoiki_linked:
    ryoikibool=[True]
else:
    ryoikibool=[True,False]

for han in np.unique(allHan):

    PapersDictSelected={k:PapersDict[k] for k in range(len(PapersDict)) if (PapersDict[k]['date']>globalmindate) & (PapersDict[k]['date']<globalmaxdate)  & (PapersDict[k]['referee'] in refbool) & (PapersDict[k]['han']==han) & (PapersDict[k]['ryoiki'] in ryoikibool)}
    TalksDictSelected={k:TalksDict[k] for k in range(len(TalksDict)) if (TalksDict[k]['date']>globalmindate) & (TalksDict[k]['date']<globalmaxdate) & (TalksDict[k]['han']==han)}
    booksDictSelected={k:booksDict[k] for k in range(len(booksDict)) if (booksDict[k]['date']>globalmindate) & (booksDict[k]['date']<globalmaxdate) & (booksDict[k]['han']==han)}
    SocialContDictSelected= {k:SocialContDict[k] for k in range(len(SocialContDict)) if (SocialContDict[k]['date']>globalmindate) & (SocialContDict[k]['date']<globalmaxdate) & (SocialContDict[k]['han']==han)}
    miscDictSelected={k:miscDict[k] for k in range(len(miscDict)) if (miscDict[k]['date']>globalmindate) & (miscDict[k]['date']<globalmaxdate) & (miscDict[k]['han']==han)}
    medDictSelected={k:medDict[k] for k in range(len(medDict)) if (medDict[k]['date']>globalmindate) & (medDict[k]['date']<globalmaxdate) & (medDict[k]['han']==han)}

    keys=list(PapersDictSelected.keys())
    datelist=[PapersDictSelected[r]['date'] for r in keys]
    arg=np.argsort(datelist)[::-1]

    document.add_paragraph(han[0]+'0'+han[1]+'班')

    countBSM=len(booksDictSelected)+len(SocialContDictSelected)+len(miscDictSelected)

    if countBSM < maxBSM:
        maxpaps=maxpap+maxBSM-countBSM
    else:
        maxpaps=maxpap
    #print(countBSM)
    textmax= ', うち'+str(min(maxpaps,len(arg)))+'件抜粋'# if inds>maxpaps else ''
    CountR=document.add_paragraph('<原著論文> 査読有計'+str(len(arg))+'件'+textmax)
    CountR.runs[0].bold=True
    inds=0
    for r in arg:
        pap=PapersDictSelected[keys[r]]
        ## to eliminate duplicates of papers
        # based on DOI
        if len(doiDict[pap['doi']]['name'])>1:
            if doiDict[pap['doi']]['count']==1:
                continue;
            titleDict[pap['papid']]['count']=1
            doiDict[pap['doi']]['count']=1
        # based on paper title
        if (len(titleDict[pap['papid']]['name'])>1):
            if titleDict[pap['papid']]['count']==1:
                continue;
            titleDict[pap['papid']]['count']=1
            doiDict[pap['doi']]['count']=1
        inds=inds+1
        if inds<=maxpaps:
            if pap['issues']:
                p = document.add_paragraph('***')
            if numberingPapers:
                if pap['ryoiki']:
                    p = document.add_paragraph(smark+str(inds)+'. ')
                else:
                    p = document.add_paragraph(str(inds)+'. ')
            else:
                if pap['ryoiki']:
                    p = document.add_paragraph(smark)
                else:
                    p = document.add_paragraph('')

            for nm in pap['authors']:
                if nm in nameListPrint:
                    listedCorrespo = any([c for c,n in zip(doiDict[pap['doi']]['Corresp'] + titleDict[pap['papid']]['Corresp'] , doiDict[pap['doi']]['name'] + titleDict[pap['papid']]['name']) if n==nm])
                    # print(nm,listedCorrespo)
                    if pap['Corresp'] | listedCorrespo:
                        p.add_run('*')
                    if daihyobuntanList[nameListPrint.index(nm)]=='D':
                        run=p.add_run()
                        run.text=nm
                        run.underline = WD_UNDERLINE.DOUBLE
                        run.font.bold =True
                    elif daihyobuntanList[nameListPrint.index(nm)]=='B':
                        run=p.add_run()
                        run.text=nm
                        run.underline = True
                        run.font.bold =True
                    else:
                        p.add_run(nm)
                else:
                    p.add_run(nm)
                p.add_run(', ')
            p.add_run(pap['text1'])
            p.add_run(pap['text2'])

    if inds != len(arg):
        textmax= ', うち'+str(min(maxpaps,inds))+'件抜粋'# if inds>maxpaps else ''
        CountR.text='<原著論文> 査読有計'+str(inds)+'件'+textmax
        CountR.runs[0].bold=True
        #replaced_text = paragraph.text.replace("before","after")

    for r in keys:
        doiDict[PapersDictSelected[r]['doi']]['count']=0
        titleDict[PapersDictSelected[r]['papid']]['count']=0


    keys=list(TalksDictSelected.keys())
    datelist=[TalksDictSelected[r]['date'] for r in keys]
    arg=np.argsort(datelist)[::-1]
    #document.add_paragraph('')
    textmax= ', うち'+str(min(maxtalk,len(arg)))+'件抜粋'# if len(arg)>maxtalk else ''
    CountG=document.add_paragraph('<学会発表・講演> 計'+str(len(arg))+'件'+textmax)
    CountG.runs[0].bold=True

    inds=0
    for r in arg:
        inds=inds+1
        if inds<=maxtalk:
            pap=TalksDictSelected[keys[r]]
            p = document.add_paragraph(str(inds)+'. ')
            nm=pap["presenter"]
            p.add_run(pap["printname"])
            p.add_run(', \"'+pap["presentation_title"]+"\"")
            p.add_run(', '+pap["event"])
            p.add_run(', '+pap["date"]+'.')

    keys=list(booksDictSelected.keys())
    datelist=[booksDictSelected[r]['date'] for r in keys]
    arg=np.argsort(datelist)[::-1]
    #document.add_paragraph('')
    CountB=document.add_paragraph('<書籍> 計'+str(len(arg))+'件')
    CountB.runs[0].bold=True
    inds=0
    for r in arg:
        inds=inds+1
        pap=booksDictSelected[keys[r]]
        p = document.add_paragraph(str(inds)+'. ')
        nm=pap['authors']
        p.add_run(pap["printname"]) 
        p.add_run(pap["book_owner_role"])
        p.add_run(pap["book_owner_range"])
        p.add_run(pap["book_title"])
        p.add_run(pap["publisher"])
        p.add_run(' '+pap["date"][:7]+'.')

    keys=list(SocialContDictSelected.keys())
    datelist=[SocialContDictSelected[r]['date'] for r in keys]
    arg=np.argsort(datelist)[::-1]
    textmax= ', うち'+str(min(maxsocial,len(arg)))+'件抜粋'# if len(arg)>maxsocial else ''
    CountO=document.add_paragraph('<アウトリーチ> 計'+str(len(arg))+'件'+textmax)
    CountO.runs[0].bold=True
    inds=0
    for r in arg:
        inds=inds+1
        if inds<=maxsocial:
            pap=SocialContDictSelected[keys[r]]
            p = document.add_paragraph(str(inds)+'. ')
            p.add_run(pap["name"])
            p.add_run(', '+pap["title"])
            p.add_run(', '+pap["event"])
            p.add_run(' '+pap["date"]+'.')

    keys=list(medDictSelected.keys())
    datelist=[medDictSelected[r]['date'] for r in keys]
    arg=np.argsort(datelist)[::-1]
    #document.add_paragraph('')
    textmax= ', うち'+str(min(maxmed,len(arg)))+'件抜粋'# if len(arg)>maxmed else ''
    CountP=document.add_paragraph('<報道> 計'+str(len(arg))+'件'+textmax)
    CountP.runs[0].bold=True
    inds=0
    for r in arg:
        inds=inds+1
        if inds<=maxsonota:
            pap=medDictSelected[keys[r]]
            p = document.add_paragraph(str(inds)+'. ')
            nm=pap['authors']
            p.add_run(pap["printname"])
            p.add_run(','+pap["paper_title"])
            p.add_run(pap["publication_name"])
            p.add_run(' '+pap["date"]+'.')

    keys=list(miscDictSelected.keys())
    datelist=[miscDictSelected[r]['date'] for r in keys]
    arg=np.argsort(datelist)[::-1]
    #document.add_paragraph('')
    textmax= ', うち'+str(min(maxsonota,len(arg)))+'件抜粋'# if len(arg)>maxsonota else ''
    CountM=document.add_paragraph('<その他> 計'+str(len(arg))+'件'+textmax)
    CountM.runs[0].bold=True
    inds=0
    for r in arg:
        inds=inds+1
        if inds<=maxsonota:
            pap=miscDictSelected[keys[r]]
            p = document.add_paragraph(str(inds)+'. ')
            nm=pap['authors']
            p.add_run(pap["printname"])
            p.add_run(','+pap["paper_title"])
            p.add_run(pap["publication_name"])
            p.add_run(' '+pap["date"]+'.')
    p = document.add_paragraph()

    p.add_run().add_break(WD_BREAK.PAGE) # page break

for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    for run in paragraph.runs:
        run.font.size = Pt(docoutputpointsize)
document.save(file_name)

# %%
# 以下「13 参考データ」入力対応
if sankodata:
    sheeturlform_csv=re.match("https://docs.google.com/spreadsheets/d/.+/",sheeturlform).group(0)+"export?format=csv"
    col_names = ['c{0:02d}'.format(i) for i in range(46)]
    name_dataform=pd.read_csv(sheeturlform_csv,header=0,names=col_names)
    name_dataform=name_dataform.sort_values('c00')
    name_dataform=name_dataform.drop_duplicates(['c02'],keep='last') # 最新のデータだけ使う
    keikakudata=name_dataform[name_dataform['c02'].isin(allkeikaku)]
    buntandata=name_dataform[~name_dataform['c02'].isin(allkeikaku)]

    timestamps=name_dataform.iloc[:,0].to_list()
    namelist=name_dataform.iloc[:,2].to_list()

    #keikakudata

# %%
if sankodata:
    FYstart=['2019-09-01','2020-04-01']
    FYend=['2020-03-31','2021-03-31']
    yearcol=['C','E']

    if 'google.colab' in str(get_ipython()):
        import urllib.request
        urllib.request.urlretrieve(blankxlsx, 'blankfile.xlsx')
        wb = openpyxl.load_workbook("blankfile.xlsx")
    else:
        wb = openpyxl.load_workbook("./inputfiles/R3中間評価報告書（1_領域全体）（13参考データExcel版）.xlsx")

    ws = wb.worksheets[0]

    document = Document()
    document.add_paragraph('カウントチェック用出力'+str(datetime.datetime.today()))

    section = document.sections[0]
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    years=2

    # 研究代表者人数
    ## R1 計画
    ws['C8'].value=len(allkeikakuPIs)

    ## R1 公募
    ws['D8'].value=0

    ## R2 計画
    ws['E8'].value=len(allkeikakuPIs)

    ## R2 公募
    ws['F8'].value=allgroupnames.count('D')

    # 研究分担者人数
    ## R1 計画
    ws['C9'].value=np.sum(keikakudata['c03'])
    ## R2 計画
    ws['E9'].value=np.sum(keikakudata['c09'])

    #研究協力者 若手研究者 外国人研究者 ポスドク RA等
    keikakulist=[col+'{0:02d}'.format(i) for col in ['C','E'] for i in range(10,15) ]
    buntanlist=[col+'{0:02d}'.format(i)  for col in ['D','F'] for i in range(10,15)]
    datapositions=['c'+'{0:02d}'.format(i) for i in [4,5,6,7,8,10,11,12,13,14]]
    for d,kei,bun in zip(datapositions,keikakulist,buntanlist):
        ws[kei].value=keikakudata[d].astype(float).sum()
        ws[bun].value=buntandata[d].astype(float).sum()

    papercelllist=[col+'{0:02d}'.format(i) for col in yearcol for i in range(21,25) ]
    yearstext=['令和元年度', '令和２年度']
    reftext=['査読あり','査読無し']
    kokunaitext=['国際誌','国内誌']
    OudanDict={}

    i=0
    for y,yt in zip(range(years),yearstext):
        for kokunai,kt in zip([False,True],kokunaitext):
            for referee,rt in zip([True,False],reftext):
                subdict={k:PapersDict[k] for k in range(len(PapersDict)) if (PapersDict[k]['date']>FYstart[y]) & (PapersDict[k]['date']<=FYend[y])  & (PapersDict[k]['referee'] == referee) & (PapersDict[k]['ryoiki'] in ryoikibool) & (PapersDict[k]['kokunai']==kokunai)}
                #print('<'+yt+' '+kt+' '+rt+'>')
                document.add_paragraph('<'+yt+' '+kt+' '+rt+'>')
                
                keys=list(subdict.keys())
                datelist=[subdict[r]['date'] for r in keys]
                arg=np.argsort(datelist)[::-1]
                inds=0
                for r in arg:
                    inds=inds+1
                    pap=subdict[keys[r]]
                    ## to eliminate duplicates of papers
                    # based on DOI
                    if len(doiDict[pap['doi']]['name'])>1:
                        if doiDict[pap['doi']]['count']==1:
                            continue;
                        titleDict[pap['papid']]['count']=1
                        doiDict[pap['doi']]['count']=1
                    # based on paper title
                    if (len(titleDict[pap['papid']]['name'])>1):
                        if titleDict[pap['papid']]['count']==1:
                            continue;
                        titleDict[pap['papid']]['count']=1
                        doiDict[pap['doi']]['count']=1
                    if pap['issues']:
                        p = document.add_paragraph('***')
                    if pap['ryoiki']:
                        p = document.add_paragraph(smark+str(inds)+'. ')
                    else:
                        p = document.add_paragraph(str(inds)+'. ')
                        p.add_run('(紐づけなし) ').font.color.rgb= RGBColor(255,0,0)
                    for nm in pap['authors']:
                        if nm in nameListPrint:
                            listedCorrespo = any([c for c,n in zip(doiDict[pap['doi']]['Corresp'] + titleDict[pap['papid']]['Corresp'] , doiDict[pap['doi']]['name'] + titleDict[pap['papid']]['name']) if n==nm])
                            # print(nm,listedCorrespo)
                            if pap['Corresp'] | listedCorrespo:
                                p.add_run('*')
                            if daihyobuntanList[nameListPrint.index(nm)]=='D':
                                p.add_run(nm).underline = WD_UNDERLINE.DOUBLE
                            elif daihyobuntanList[nameListPrint.index(nm)]=='B':
                                p.add_run(nm).underline = True
                            else:
                                p.add_run(nm)
                        else:
                            p.add_run(nm)
                        p.add_run(', ')
                    p.add_run(pap['text1'])
                    p.add_run(pap['text2'])
                    if pap['oudan']:
                        #print(pap['oudan'])
                        p.add_run(' ('+ pap['oudan']+')').font.color.rgb = RGBColor(0,0,255)
                        if pap['oudan'] in OudanDict.keys():
                            OudanDict[pap['oudan']]['count']+=1
                        else:
                            OudanDict[pap['oudan']]={}
                            OudanDict[pap['oudan']]['count']=1
                p=document.add_paragraph('...(計')
                p.add_run(str(inds)).font.color.rgb= RGBColor(255,0,0)
                p.add_run('件)')
                document.add_paragraph('')
                ws[papercelllist[i]]=inds
                i=i+1

    p=document.add_paragraph('')
    p.add_run().add_break(WD_BREAK.PAGE) # page break

    document.add_paragraph('共同研究リスト')
    DKtext=['大学','企業・公共団体']
    kokunaikaigaitext=['国内','海外']

    inputpositionsKyodo=['c'+'{0:02d}'.format(i) for i in range(24,32)]
    outputpositionsKyodo=[col+'{0:02d}'.format(i) for col in ['C','E'] for i in range(47,51) ]
    for iK,oK in zip(inputpositionsKyodo,outputpositionsKyodo):
        ws[oK].value=name_dataform[iK].astype(float).sum()

    inputpositionsKyodo=['c'+'{0:02d}'.format(i) for i in range(32,40)]
    outputpositionsKyodo=[col+'{0:02d}'.format(i) for col in ['D','F'] for i in range(47,51) ]
    printtext=[yt + ' '+ kt +' '+ dk + ' '+ '契約書なし' for yt in yearstext  for dk in DKtext for kt in kokunaikaigaitext]

    for iK,oK,ptxt in zip(inputpositionsKyodo,outputpositionsKyodo,printtext):
        sumKyodo=0
        document.add_paragraph(ptxt)
        for a,piname in zip(name_dataform[iK],name_dataform['c02']):
            if (type(a)==str): 
                if (len(a)>4): # なし、0などを省くために5文字以上
                    b=re.split('[、。,]', a)
                    sumKyodo+=len(b)
                    p =document.add_paragraph(piname+' 研: ')
                    for bb in b:
                        p.add_run(bb).underline = True
                        p.add_run('　')

        ws[oK].value=sumKyodo
        p=document.add_paragraph('...(計')
        p.add_run(str(sumKyodo)).font.color.rgb= RGBColor(255,0,0)
        p.add_run('件)')
        document.add_paragraph('')

    p=document.add_paragraph('')
    p.add_run().add_break(WD_BREAK.PAGE) # page break


    document.add_paragraph('受賞リスト')

    keys=list(AwardsDict.keys())
    datelist=[AwardsDict[r]['award_date'] for r in keys]
    arg=np.argsort(datelist)[::-1]
    inds=0
    for r in arg:
        if (AwardsDict[keys[r]]['award_date']>globalmindate) & (AwardsDict[keys[r]]['award_date']<globalmaxdate):
            inds=inds+1
            p = document.add_paragraph(str(inds)+'. ')
            p.add_run(AwardsDict[keys[r]]['name'])
            p.add_run(', '+AwardsDict[keys[r]]['award_name'])
            p.add_run(', '+AwardsDict[keys[r]]['association'])
            p.add_run(', '+AwardsDict[keys[r]]['award_date']+'.')


    p=document.add_paragraph('')
    p.add_run().add_break(WD_BREAK.PAGE) # page break

    document.add_paragraph('国際会議招待講演リスト')
    keys=list(TalksDict.keys())
    datelist=[TalksDict[r]['date'] for r in keys]
    arg=np.argsort(datelist)[::-1]
    inds=0
    for r in arg:
        if TalksDict[keys[r]]["international"] & (TalksDict[keys[r]]['date']>globalmindate) & (TalksDict[keys[r]]['date']<globalmaxdate) & (TalksDict[keys[r]]["invited"] | (TalksDict[keys[r]]["keyoral"] in ['invited_oral_presentation','keynote_oral_presentation','nominated_symposium'])):
            inds=inds+1
            pap=TalksDict[keys[r]]
            p = document.add_paragraph(str(inds)+'. ')
            nm=pap["presenter"]
            p.add_run(pap["printname"])
            p.add_run(', \"'+pap["presentation_title"]+"\"")
            p.add_run(' ('+pap["keyoral"]+')').font.color.rgb= RGBColor(255,0,0)
            p.add_run(', '+pap["event"])
            p.add_run(', '+pap["date"]+'.')


    p=document.add_paragraph('')
    p.add_run().add_break(WD_BREAK.PAGE) # page break

    document.add_paragraph('アウトリーチリスト')
    keys=list(SocialContDict.keys())
    datelist=[SocialContDict[r]['date'] for r in keys]
    arg=np.argsort(datelist)[::-1]

    inds=0
    for r in arg:
        pap=SocialContDict[keys[r]]
        if (SocialContDict[keys[r]]['date']>globalmindate) & (SocialContDict[keys[r]]['date']<globalmaxdate):
            inds=inds+1
            pap=SocialContDict[keys[r]]
            p = document.add_paragraph(str(inds)+'. ')
            p.add_run(pap["name"])
            p.add_run(', '+pap["title"])
            p.add_run(', '+pap["event"])
            p.add_run(' '+pap["date"]+'.')
            
    p=document.add_paragraph('')
    p.add_run().add_break(WD_BREAK.PAGE) # page break

    document.add_paragraph('報道リスト')
    keys=list(medDict.keys())
    datelist=[medDict[r]['date'] for r in keys]
    arg=np.argsort(datelist)[::-1]

    inds=0
    for r in arg:
        if (medDict[keys[r]]['date']>globalmindate) & (medDict[keys[r]]['date']<globalmaxdate):
            inds=inds+1
            pap=medDict[keys[r]]
            p = document.add_paragraph(str(inds)+'. ')
            p.add_run(pap["printname"])
            p.add_run(','+pap["paper_title"])
            p.add_run(pap["publication_name"])
            p.add_run(' '+pap["date"]+'.')

    p=document.add_paragraph('')
    p.add_run().add_break(WD_BREAK.PAGE) # page break

    document.add_paragraph('就職状況リスト')

    for s in name_dataform['c22']:
        if type(s) is str:
            p = document.add_paragraph(s)

    wb.save(file_name_xlsx)
    document.save(file_name_check)


# %%
globalmindate

# %%
if ('google.colab' in str(get_ipython())):
    files.download(file_name)
    if sankodata:
        files.download(file_name_xlsx)
        files.download(file_name_check)


