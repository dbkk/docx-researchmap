# %%
## いじるのはこのセルのパラメータのみでOK
file_name_download = "IPB-20210126.docx" #ダウンロードされるファイル名
globalmindate='2020-04-01' #これより後の業績を集める
globalmaxdate='2021-04-01' #これより前の業績を集める
smark='' #謝辞に課題番号ありの論文にマーク付ける場合はここで指定。
allenglish = True #名前表記をすべて英語で統一する場合はTrue, 論文以外の名前表記を日本語にする場合False
numberingPapers = True #出力の際に論文をナンバリング
peer_reviewed = False #査読ありのチェックが入った論文だけに限定する場合はTrue
sheeturl='https://docs.google.com/spreadsheets/d/1wce1XHSFGSBttupnSIqe_5abtijBb_hBYM2bfaV9Jn4/edit#gid=0' #作成したgoogle spreadsheetのアドレス

# %%
import requests,json,sys,os, gspread, time, re
import numpy as np
if 'google.colab' in str(get_ipython()):
    %pip install python-docx
    from google.colab import files,auth
    from oauth2client.client import GoogleCredentials
    outputdirectory = ''
else:
    outputdirectory = '../docx-researchmap-outputs/' #ローカルで実行する場合は保存ファイルのディレクトリを適当に指定
    os.makedirs(outputdirectory,exist_ok=True)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
import pandas as pd
file_name=outputdirectory+file_name_download

# %%
#スプレッドシートをダウンロード
sheeturl_csv=re.match("https://docs.google.com/spreadsheets/d/.+/",sheeturl).group(0)+"export?format=csv"
name_data=pd.read_csv(sheeturl_csv)
name_data

# %%
membernum=len(name_data)

allnames=(name_data["Surname"]+' '+name_data["First name"]).to_list()
allSurname=name_data["Surname"].to_list()
allnamesJP=(name_data["苗字"]+" "+name_data["名"]).to_list()
allmembers=name_data["researchmapID"].to_list()
allDB=name_data["代表分担協力"].values
allDaihyoBuntan=list(allDB)
grant_numbers=name_data["grantID"].to_list()
allmindate=name_data["Start date"].to_list()
allmaxdate=name_data["End date"].to_list()

#Exception names handling
altname2,altname3=name_data['著者名（2個目）'],name_data['著者名（3個目）']
arraltname2,arraltname3=altname2.values,altname3.values
nameList=allnames+list(arraltname2[~(pd.isna(altname2).values)])+list(arraltname3[~(pd.isna(altname3).values)])
nameList = [n.strip() for n in nameList]
daihyobuntanList=allDaihyoBuntan+list(allDB[~(pd.isna(altname2).values)])+list(allDB[~(pd.isna(altname3).values)])

# %%
# Function to set the name order (surname-firstname)
def SurnameFirst(namesDic,sn):
    oldnamelist=[]
    swap=0
    for indiv in namesDic:
        oldnamelist=oldnamelist+[indiv['name'].replace(',','').replace('.','')]
        #print(oldnamelist)
    for name in oldnamelist:
        if sn in name.split(' '):
            if name.split(' ').index(sn)==0:
                swap=0
                break;
            else:
                swap=1
                break;
    if swap:
        newnamelist=[]
        for name in oldnamelist:
            namesplit=name.split(' ')
            names=[namesplit[-1]]+namesplit[:-1]
            newnamelist=newnamelist+[' '.join(names)]
    else:
        newnamelist=oldnamelist
    return newnamelist

def ReturnDictWOerror(dictdata,key,nodata):
    if key in dictdata.keys():
        return dictdata[key]
    else:
        return nodata

def ReturnDictContent(dictdata,key,key1,nodata=''):
    d=ReturnDictWOerror(dictdata,key,nodata)
    d1=ReturnDictWOerror(dictdata,key1,nodata)
    if d!=nodata:
        return d
    else:
        return d1

def commaR(vol,spage):
    if (vol=='') & (spage==''):
        return ''
    elif (vol=='') | (spage==''):
        return ' '
    else:
        return ', '

# %%
#download json files from researchmap
url = "https://api.researchmap.jp/"
itemslist = ["published_papers","research_projects","misc","presentations","books_etc"]
jsonfiles={}
for name in allmembers:
  print('downloading: '+name)
  jsonfiles[name]={}
  for it in itemslist:
    r1 = requests.get(url+name+'/'+it)
    jsonfiles[name][it]=json.loads(r1.text)
    if 'error' in jsonfiles[name][it].keys():
      print(jsonfiles[name][it]['error'])
      print("  error in:"+it)

# %%
# make dictionary of all papers
i=0
PapersDict={}

doilist=[]
doiDict={}
titlelist=[]
titleDict={}
#dc = Document()
for ids,fullname,dh,mindate,maxdate in zip(allmembers,allnames,allDaihyoBuntan,allmindate,allmaxdate):
    surname=fullname.split(' ')[0]
    dfP = jsonfiles[ids]["published_papers"]
    dfG = jsonfiles[ids]["research_projects"]
    if 'items' in dfG.keys():
        grantID="0"
        for dfs in dfG['items']:
            if 'identifiers' in dfs.keys():
                if 'grant_number' in dfs['identifiers'].keys():
                    if dfs['identifiers']['grant_number'][0] in grant_numbers:
                        grantID=dfs['rm:id']
                        break
    if 'items' in dfP.keys():    
        for dfs in dfP['items']:
            if "authors" not in dfs.keys():
                continue
            if ('identifiers' in dfs.keys()) & (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
                doinum=[0]
                if 'doi' in dfs['identifiers'].keys():
                    doinum=dfs['identifiers']['doi']

                PapersDict[i]={}
                PapersDict[i]['issues']=False
                PapersDict[i]['preprint']=False
                correspo=False
                Ryoiki=False
                if 'research_project_id' in dfs['identifiers'].keys():
                    if grantID in dfs['identifiers']['research_project_id']:
                        Ryoiki=True
                if "published_paper_owner_roles" in dfs.keys():
                    if "corresponding" in dfs["published_paper_owner_roles"]:
                        correspo=True

                jname=''        
                if "publication_name" in dfs.keys():
                    jname=ReturnDictContent(dfs["publication_name"],'en','ja','').upper()

                if jname =='ARXIV':
                    PapersDict[i]['preprint']=True
                    if "arxiv_id" in dfs['identifiers'].keys():
                        jname=dfs['identifiers']['arxiv_id'][0] + ' (preprint)'
                    else:
                        jname='arxiv'
                
                if not("publication_name" in dfs.keys()):
                    if "arxiv_id" in dfs['identifiers'].keys():
                        jname=dfs['identifiers']['arxiv_id'][0] + ' (preprint)'
                        PapersDict[i]['preprint']=True
                    elif doinum[0]!=0:
                        jname='DOI: '+doinum[0]
                        PapersDict[i]['preprint']=True
                    else:
                        jname='journal unspecified'
                        PapersDict[i]['issues']=True
                    
                Sname=SurnameFirst(ReturnDictContent(dfs["authors"],'en','ja',''),surname)

                spage=''
                if "starting_page" in dfs.keys():
                    if dfs["starting_page"]!='':
                        spage=dfs["starting_page"]

                vol=''
                if "volume" in dfs.keys():
                    if dfs["volume"]!='':
                        vol=' '+dfs["volume"]
                if doinum in doilist:
                    doiDict[doinum[0]]['name']=doiDict[doinum[0]]['name']+[fullname]
                    doiDict[doinum[0]]['Corresp']=doiDict[doinum[0]]['Corresp']+[correspo]
                else:
                    doiDict[doinum[0]]={}
                    doiDict[doinum[0]]['name']=[fullname]
                    doiDict[doinum[0]]['Corresp']=[correspo]
                    doiDict[doinum[0]]['count']=0
                    doilist=doilist+[doinum[0]]
                
                papertitle=ReturnDictContent(dfs['paper_title'],'en','ja','')
                papid=papertitle.upper().rstrip('.')

                if papid in titlelist:
                    titleDict[papid]['name'] = titleDict[papid]['name']+[fullname]
                    titleDict[papid]['Corresp'] = titleDict[papid]['Corresp']+[correspo]                    
                else:
                    titlelist = titlelist + [papid]
                    titleDict[papid] = {}
                    titleDict[papid]['name'] = [fullname]
                    titleDict[papid]['Corresp'] = [correspo]
                    titleDict[papid]['count']=0

                text1="\""+papertitle+"\"" +', '
                text2=jname+','+vol+commaR(vol,spage)+spage+ ' ('+dfs["publication_date"][:4] +').'
                PapersDict[i].update({
                    'text1': text1,
                    'text2': text2,
                    'papid': papid,
                    'researcher': fullname,
                    'authors': Sname,
                    'date': dfs["publication_date"],
                    'referee': ReturnDictContent(dfs, 'referee', 'referee', False),
                    'doi': doinum[0],
                    'ryoiki': Ryoiki,
                    'Daihyo': dh,
                    'Corresp': correspo
                })
                i=i+1

# %%
# make dictionary of all talks
TalksDict={}
i=0
for ids,fullname,fullnameJP,dh,mindate,maxdate in zip(allmembers,allnames,allnamesJP,allDaihyoBuntan,allmindate,allmaxdate):
    dfPr = jsonfiles[ids]["presentations"]
    if 'items' in dfPr.keys():
        for dfs in dfPr['items']:
            if all([a in dfs.keys() for a in ['invited',"presentation_title","event",'publication_date',"presenters"]]):
                if (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
                    if dfs['invited']:
                        if ('en' in dfs["presenters"].keys()):
                            pname=dfs["presenters"]["en"][0]["name"]
                        else:
                            pname=dfs["presenters"]["ja"][0]["name"]
                        ename=ReturnDictContent(dfs["event"],'en','ja','')
                        ptitle=ReturnDictContent(dfs["presentation_title"],'en','ja','')
                        pdate=dfs["publication_date"]
                        TalksDict[i] = {
                            'presenter': fullname,
                            'printname': fullname if allenglish else fullnameJP,
                            'event': ename,
                            'presentation_title': ptitle,
                            'date': pdate
                        }

                        i=i+1

# %%
# make dictionary of all books_etc
booksDict={}
i=0
for ids,fullname,fullnameJP,dh,mindate,maxdate in zip(allmembers,allnames,allnamesJP,allDaihyoBuntan,allmindate,allmaxdate):
  dfM = jsonfiles[ids]["books_etc"]
  if 'items' in dfM.keys():
    for dfs in dfM['items']:
      if all([a in dfs.keys() for a in ['authors',"book_title","publication_date"]]):
        if (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
            pname = dfs["authors"]["ja"][0]["name"] if "ja" in dfs["authors"] else dfs["authors"]["en"][0]["name"]
            ename = ReturnDictContent(dfs["book_title"], "ja", "en", "")
            eoname = " \'" + ReturnDictContent(dfs.get("book_owner_range", ""), "ja", "en", "") + "\'," if dfs.get("book_owner_range") else ""
            brole = f" ({dfs['book_owner_role']})" if "book_owner_role" in dfs else ","
            pub = f" {ReturnDictContent(dfs.get('publisher', ''), 'ja', 'en', '')}," if "publisher" in dfs else ""
            pdate=dfs["publication_date"]
            booksDict[i] = {
                "authors": fullname,
                "printname": fullname if allenglish else fullnameJP,
                "book_title": f' "{ename}",',
                "book_owner_role": brole,
                "book_owner_range": eoname,
                "publisher": pub,
                "date": pdate
            }

            i=i+1

# %%
# make dictionary of all MISCs
miscDict={}
i=0
for ids,fullname,fullnameJP,dh,mindate,maxdate in zip(allmembers,allnames,allnamesJP,allDaihyoBuntan,allmindate,allmaxdate):
    dfM = jsonfiles[ids]["misc"]
    if 'items' in dfM.keys():
        for dfs in dfM['items']:
            if all([a in dfs.keys() for a in ['authors',"paper_title","publication_date","publication_name"]]):
                if  (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
                    authors = dfs['authors'].get('ja', []) or dfs['authors'].get('en', [])
                    fullname = authors[0]['name']
                    allenglish = True  # assuming this variable is already set before this code snippet

                    ename = ReturnDictContent(dfs['paper_title'], 'ja', 'en')
                    ptitle = ReturnDictContent(dfs['publication_name'], 'ja', 'en')
                    pdate = dfs['publication_date']

                    miscDict[i] = {
                        'authors': fullname,
                        'printname': fullname if allenglish else fullnameJP,
                        'paper_title': f" '{ename}',",
                        'publication_name': f' "{ptitle}",',
                        'date': pdate,
                        }
                    i=i+1

# %%
# generate docx

if peer_reviewed:
    PapersDictSelected={k:PapersDict[k] for k in range(len(PapersDict)) if (PapersDict[k]['date']>globalmindate) & (PapersDict[k]['date']<globalmaxdate)  & (PapersDict[k]['referee'])}
else:
    PapersDictSelected={k:PapersDict[k] for k in range(len(PapersDict)) if (PapersDict[k]['date']>globalmindate) & (PapersDict[k]['date']<globalmaxdate)}   

keys=list(PapersDictSelected.keys())
datelist=[PapersDictSelected[r]['date'] for r in keys]
arg=np.argsort(datelist)[::-1]

document = Document()
document.add_paragraph('原著論文')
inds=0
for r in arg:
    inds=inds+1
    pap=PapersDictSelected[keys[r]]
    ## to eliminate duplicates of papers
    # based on DOI
    if len(doiDict[pap['doi']]['name'])>1:
        if doiDict[pap['doi']]['count']==1:
            continue;
        titleDict[pap['papid']]['count']=1
        doiDict[pap['doi']]['count']=1
    # based on paper title
    if (len(titleDict[pap['papid']]['name'])>1):
        if titleDict[pap['papid']]['count']==1:
            continue;
        titleDict[pap['papid']]['count']=1
        doiDict[pap['doi']]['count']=1

    if pap['issues']:
        p = document.add_paragraph('***')

    if numberingPapers:
        if pap['ryoiki']:
            p = document.add_paragraph(smark+str(inds)+'. '+pap['text1'])
        else:
            p = document.add_paragraph(str(inds)+'. '+pap['text1'])
    else:
        if pap['ryoiki']:
            p = document.add_paragraph(smark+pap['text1'])
        else:
            p = document.add_paragraph(pap['text1'])
    for nm in pap['authors']:
        if nm in nameList:
            listedCorrespo = any([c for c,n in zip(doiDict[pap['doi']]['Corresp'] + titleDict[pap['papid']]['Corresp'] , doiDict[pap['doi']]['name'] + titleDict[pap['papid']]['name']) if n==nm])
            # print(nm,listedCorrespo)
            if pap['Corresp'] | listedCorrespo:
                p.add_run('*')
            if daihyobuntanList[nameList.index(nm)]=='D':
                p.add_run(nm).underline = WD_UNDERLINE.DOUBLE
            elif daihyobuntanList[nameList.index(nm)]=='B':
                p.add_run(nm).underline = True
            else:
                p.add_run(nm)
        else:
            p.add_run(nm)
        p.add_run(', ')
    p.add_run(pap['text2'])

for r in keys:
    doiDict[PapersDictSelected[r]['doi']]['count']=0
    titleDict[PapersDictSelected[r]['papid']]['count']=0

TalksDictSelected={k:TalksDict[k] for k in range(len(TalksDict)) if (TalksDict[k]['date']>globalmindate) & (TalksDict[k]['date']<globalmaxdate) }

keys=list(TalksDictSelected.keys())
datelist=[TalksDictSelected[r]['date'] for r in keys]
arg=np.argsort(datelist)[::-1]
document.add_paragraph('')
document.add_paragraph('学会発表・講演（招待あり）')
inds=0
for r in arg:
    inds=inds+1
    pap=TalksDictSelected[keys[r]]
    p = document.add_paragraph(str(inds)+'. ')
    nm=pap["presenter"]
    if daihyobuntanList[nameList.index(nm)]=='D':
        p.add_run(pap["printname"]).underline = WD_UNDERLINE.DOUBLE
    elif daihyobuntanList[nameList.index(nm)]=='B':
        p.add_run(pap["printname"]).underline = True
    else:
        p.add_run(pap["printname"])
    p.add_run(', \"'+pap["presentation_title"]+"\"")
    p.add_run(', '+pap["event"])
    p.add_run(', '+pap["date"]+'.')

booksDictSelected={k:booksDict[k] for k in range(len(booksDict)) if (booksDict[k]['date']>globalmindate) & (booksDict[k]['date']<globalmaxdate) }

keys=list(booksDictSelected.keys())
datelist=[booksDictSelected[r]['date'] for r in keys]
arg=np.argsort(datelist)[::-1]
document.add_paragraph('')
document.add_paragraph('書籍')
inds=0
for r in arg:
    inds=inds+1
    pap=booksDictSelected[keys[r]]
    p = document.add_paragraph(str(inds)+'. ')
    nm=pap['authors']
    ## if underlines are required
    # if daihyobuntanList[nameList.index(nm)]=='D':
    #     p.add_run(pap["printname"]).underline = WD_UNDERLINE.DOUBLE
    # elif daihyobuntanList[nameList.index(nm)]=='B':
    #     p.add_run(pap["printname"]).underline = True
    # else:
    #     p.add_run(pap["printname"])
    p.add_run(pap["printname"]) 
    p.add_run(pap["book_owner_role"])
    p.add_run(pap["book_owner_range"])
    p.add_run(pap["book_title"])
    p.add_run(pap["publisher"])
    p.add_run(' '+pap["date"][:7]+'.')

miscDictSelected={k:miscDict[k] for k in range(len(miscDict)) if (miscDict[k]['date']>globalmindate) & (miscDict[k]['date']<globalmaxdate) }

keys=list(miscDictSelected.keys())
datelist=[miscDictSelected[r]['date'] for r in keys]
arg=np.argsort(datelist)[::-1]
document.add_paragraph('')
document.add_paragraph('その他')
inds=0
for r in arg:
    inds=inds+1
    pap=miscDictSelected[keys[r]]
    p = document.add_paragraph(str(inds)+'. ')
    nm=pap['authors']
    if daihyobuntanList[nameList.index(nm)]=='D':
        p.add_run(pap["printname"]).underline = WD_UNDERLINE.DOUBLE
    elif daihyobuntanList[nameList.index(nm)]=='B':
        p.add_run(pap["printname"]).underline = True
    else:
        p.add_run(pap["printname"])
    p.add_run(','+pap["paper_title"])
    p.add_run(pap["publication_name"])
    p.add_run(' '+pap["date"]+'.')
document.save(file_name)

# %%
if 'google.colab' in str(get_ipython()):
    files.download(file_name)


