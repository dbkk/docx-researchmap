# %%
## いじるのはこのセルのパラメータのみでOK
file_name_download = "kojin-20210530.docx" #ダウンロードされるファイル名
mindate='2020-04-01' #これより後の業績を集める
maxdate='2021-04-01' #これより前の業績を集める
smark='▲' #謝辞に課題番号ありの論文にマーク付ける場合はここで指定。中間報告では'▲'。
allenglish = True #名前表記をすべて英語で統一する場合はTrue, 論文以外の名前表記を日本語にする場合False
numberingPapers = True #出力の際に論文をナンバリング
peer_reviewed = False #査読ありのチェックが入った論文だけに限定する場合はTrue
sheeturl='https://docs.google.com/spreadsheets/d/1hf2oZbtyu-jCxEiljA8KiUVLZvIybFf0BrfnLc4nuYA/edit?usp=sharing' #作成したgoogle spreadsheetのアドレス
SNfirst = False #英語名前表記をすべて名字先で統一する場合はTrue, 名字後で統一する場合はFalse
doiadd = True #

# %%
import requests,json,sys,os, gspread, time, re
import numpy as np
if 'google.colab' in str(get_ipython()):
    %pip install python-docx
    from google.colab import files,auth
    from oauth2client.client import GoogleCredentials
    outputdirectory = ''
else:
    outputdirectory = '../docx-researchmap-outputs/' #ローカルで実行する場合は保存ファイルのディレクトリを適当に指定
    os.makedirs(outputdirectory,exist_ok=True)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
import pandas as pd
file_name=outputdirectory+file_name_download

# %%
#スプレッドシートをダウンロード
sheeturl_csv=re.match("https://docs.google.com/spreadsheets/d/.+/",sheeturl).group(0)+"export?format=csv"
name_data=pd.read_csv(sheeturl_csv)

# %%
membernum=len(name_data)

allnames=(name_data["Surname"]+' '+name_data["First name"]).to_list()
allSurname=name_data["Surname"].to_list()
allnamesJP=(name_data["苗字"]+" "+name_data["名"]).to_list()
allmembers=name_data["researchmapID"].to_list()
allDaihyoBuntan=name_data["代表分担協力"].to_list()
grant_numbers=name_data["grantID"].to_list()

#Exception handling
nameList=allnames+['Sagawa T', 'Tetsuya J kobayashi']
daihyobuntanList=allDaihyoBuntan+['D','D']

# %%
# Function to set the name order (surname-firstname)
def SurnameFirst(namesDic,sn):
    oldnamelist=[]
    swap=0
    for indiv in namesDic:
        oldnamelist=oldnamelist+[indiv['name'].replace(',','').replace('.','')]
        #print(oldnamelist)
    for name in oldnamelist:
        if sn in name.split(' '):
            if name.split(' ').index(sn)==0: # surname first
                swap= True ^ SNfirst
                break;
            else:
                swap= False ^ SNfirst
                break;
    if swap:
        newnamelist=[]
        for name in oldnamelist:
            namesplit=name.split(' ')
            names=[namesplit[-1]]+namesplit[:-1]
            newnamelist=newnamelist+[' '.join(names)]
    else:
        newnamelist=oldnamelist
    return newnamelist

def ReturnDictWOerror(dictdata,key,nodata):
    if key in dictdata.keys():
        return dictdata[key]
    else:
        return nodata

def ReturnDictContent(dictdata,key,key1,nodata=''):
    d=ReturnDictWOerror(dictdata,key,nodata)
    d1=ReturnDictWOerror(dictdata,key1,nodata)
    if d!=nodata:
        return d
    else:
        return d1

def commaR(vol,spage):
    if (vol=='') & (spage==''):
        return ''
    elif (vol=='') | (spage==''):
        return ' '
    else:
        return ', '

# %%
#download json files from researchmap
url = "https://api.researchmap.jp/"
itemslist = ["published_papers","research_projects","misc","presentations","books_etc"]
jsonfiles={}
for name in allmembers:
  jsonfiles[name]={}
  for it in itemslist:
    r1 = requests.get(url+name+'/'+it)
    jsonfiles[name][it]=json.loads(r1.text)

# %%
# make dictionary of all papers
i=0
PapersDict={}

doilist=[]
doiDict={}
#dc = Document()
for ids,fullname,dh in zip(allmembers,allnames,allDaihyoBuntan):
    surname=fullname.split(' ')[0]
    dfP = jsonfiles[ids]["published_papers"]
    dfG = jsonfiles[ids]["research_projects"]
    if 'items' in dfG.keys():
        grantID="0"
        for dfs in dfG['items']:
            if 'identifiers' in dfs.keys():
                if 'grant_number' in dfs['identifiers'].keys():
                    if dfs['identifiers']['grant_number'][0] in grant_numbers:
                        grantID=dfs['rm:id']
                        break
    if 'items' in dfP.keys():    
        for dfs in dfP['items']:
            if "authors" not in dfs.keys():
                continue
            if 'identifiers' in dfs.keys():
                doinum=[0]
                if 'doi' in dfs['identifiers'].keys():
                    doinum=dfs['identifiers']['doi']

                PapersDict[i]={}
                correspo=False
                Ryoiki=False
                if 'research_project_id' in dfs['identifiers'].keys():
                    if grantID in dfs['identifiers']['research_project_id']:
                        Ryoiki=True
                if "published_paper_owner_roles" in dfs.keys():
                    if "corresponding" in dfs["published_paper_owner_roles"]:
                        correspo=True

                jname=''        
                if "publication_name" in dfs.keys():
                    jname=ReturnDictContent(dfs["publication_name"],'en','ja','').upper()

                if jname =='ARXIV':
                    PapersDict[i]['preprint']=True
                    if "arxiv_id" in dfs['identifiers'].keys():
                        jname=dfs['identifiers']['arxiv_id'][0] + ' (preprint)'
                    else:
                        jname='arxiv'
                
                if not("publication_name" in dfs.keys()):
                    if "arxiv_id" in dfs['identifiers'].keys():
                        jname=dfs['identifiers']['arxiv_id'][0] + ' (preprint)'
                        PapersDict[i]['preprint']=True
                    elif doinum[0]!=0:
                        jname='DOI: '+doinum[0]
                        PapersDict[i]['preprint']=True
                    else:
                        jname='journal unspecified'
                        PapersDict[i]['issues']=True
                    
                Sname=SurnameFirst(ReturnDictContent(dfs["authors"],'en','ja',''),surname)

                spage=''
                if "starting_page" in dfs.keys():
                    if dfs["starting_page"]!='':
                        spage=dfs["starting_page"]

                vol=''
                if "volume" in dfs.keys():
                    if dfs["volume"]!='':
                        vol=' '+dfs["volume"]
                if doinum in doilist:
                    doiDict[doinum[0]]['name']=doiDict[doinum[0]]['name']+[fullname]
                    doiDict[doinum[0]]['Corresp']=doiDict[doinum[0]]['Corresp']+[correspo]
                else:
                    doiDict[doinum[0]]={}
                    doiDict[doinum[0]]['name']=[fullname]
                    doiDict[doinum[0]]['Corresp']=[correspo]
                    doilist=doilist+[doinum[0]]
                text1="\""+ReturnDictContent(dfs['paper_title'],'en','ja','')+"\"" +', '
                text2=jname+','+vol+commaR(vol,spage)+spage+ ' ('+dfs["publication_date"][:4] +').'
                PapersDict[i]['text1']=text1
                PapersDict[i]['text2']=text2
                PapersDict[i]['researcher']=fullname
                PapersDict[i]['authors']=Sname
                PapersDict[i]['date']=dfs["publication_date"]
                PapersDict[i]['referee']=ReturnDictContent(dfs,'referee','referee',False)
                PapersDict[i]['doi']=doinum[0]
                PapersDict[i]['ryoiki']=Ryoiki
                PapersDict[i]['Daihyo']=dh
                PapersDict[i]['Corresp']=correspo
                i=i+1

# %%
# make dictionary of all talks
TalksDict={}
i=0
for ids,fullname,fullnameJP,dh in zip(allmembers,allnames,allnamesJP,allDaihyoBuntan):
    dfPr = jsonfiles[ids]["presentations"]
    if 'items' in dfPr.keys():
        for dfs in dfPr['items']:
            if all([a in dfs.keys() for a in ['invited',"presentation_title","event",'publication_date',"presenters"]]):
                if dfs['invited']:
                    if ('en' in dfs["presenters"].keys()):
                        pname=dfs["presenters"]["en"][0]["name"]
                    else:
                        pname=dfs["presenters"]["ja"][0]["name"]
                    ename=ReturnDictContent(dfs["event"],'en','ja','')
                    ptitle=ReturnDictContent(dfs["presentation_title"],'en','ja','')
                    pdate=dfs["publication_date"]
                    TalksDict[i]={}
                    TalksDict[i]["presenter"]=fullname
                    if allenglish:
                        TalksDict[i]['printname']=fullname
                    else:
                        TalksDict[i]["printname"]=fullnameJP
                    TalksDict[i]["event"]=ename
                    TalksDict[i]["presentation_title"]=ptitle
                    TalksDict[i]["date"]=pdate
                    i=i+1

# %%
# make dictionary of all books_etc
booksDict={}
i=0
for ids,fullname,fullnameJP,dh in zip(allmembers,allnames,allnamesJP,allDaihyoBuntan):
  dfM = jsonfiles[ids]["books_etc"]
  if 'items' in dfM.keys():
    for dfs in dfM['items']:
      if all([a in dfs.keys() for a in ['authors',"book_title","publication_date"]]):
        if (dfs["publication_date"]>=mindate) & (dfs["publication_date"]<=maxdate):
          if ('ja' in dfs["authors"].keys()):
              pname=dfs["authors"]["ja"][0]["name"]
          else:
              pname=dfs["authors"]["en"][0]["name"]
          ename=ReturnDictContent(dfs["book_title"],'ja','en','')
          if "book_owner_range" in dfs.keys():
            eoname=" \'"+ReturnDictContent(dfs["book_owner_range"],'ja','en','')+"\',"
          else:
            eoname=''
          if "book_owner_role" in dfs.keys():
            brole=" ("+dfs["book_owner_role"]+"),"
          else:
            brole=','
          if "publisher" in dfs.keys():
            pub=" "+ReturnDictContent(dfs["publisher"],'ja','en','')+","
          else:
            pub=''
          pdate=dfs["publication_date"]
          booksDict[i]={}
          booksDict[i]['authors']=fullname
          if allenglish:
            booksDict[i]['printname']=fullname
          else:
            booksDict[i]['printname']=fullnameJP
          booksDict[i]["book_title"]=' \"'+ename+'\",'
          booksDict[i]["book_owner_role"]=brole
          booksDict[i]["book_owner_range"]=eoname
          booksDict[i]["publisher"]=pub
          booksDict[i]["date"]=pdate
          i=i+1

# %%
# make dictionary of all MISCs
miscDict={}
i=0
for ids,fullname,fullnameJP,dh in zip(allmembers,allnames,allnamesJP,allDaihyoBuntan):
  dfM = jsonfiles[ids]["misc"]
  if 'items' in dfM.keys():
    for dfs in dfM['items']:
      if all([a in dfs.keys() for a in ['authors',"paper_title","publication_date","publication_name"]]):
        if ('ja' in dfs["authors"].keys()):
            pname=dfs["authors"]["ja"][0]["name"]
        else:
            pname=dfs["authors"]["en"][0]["name"]
        if ('ja' in dfs["paper_title"].keys()):
            ename=dfs["paper_title"]["ja"]
        else:
            ename=dfs["paper_title"]["en"]
        if ('ja' in dfs["publication_name"].keys()):
            ptitle=dfs["publication_name"]["ja"]
        else:
            ptitle=dfs["publication_name"]["en"]
        pdate=dfs["publication_date"]
        #print(pname,ename,pdate)
        miscDict[i]={}
        miscDict[i]['authors']=fullname
        if allenglish:
          miscDict[i]['printname']=fullname
        else:
          miscDict[i]['printname']=fullnameJP

        miscDict[i]["paper_title"]=' \''+ename+'\','
        miscDict[i]["publication_name"]=' \"'+ptitle+'\",'
        miscDict[i]["date"]=pdate
        i=i+1

# %%
# generate docx

if peer_reviewed:
    PapersDictSelected={k:PapersDict[k] for k in range(len(PapersDict)) if (PapersDict[k]['date']>mindate) & (PapersDict[k]['date']<maxdate)  & (PapersDict[k]['referee'])}
else:
    PapersDictSelected={k:PapersDict[k] for k in range(len(PapersDict)) if (PapersDict[k]['date']>mindate) & (PapersDict[k]['date']<maxdate) }


keys=list(PapersDictSelected.keys())
datelist=[PapersDictSelected[r]['date'] for r in keys]
arg=np.argsort(datelist)[::-1]

document = Document()
document.add_paragraph('原著論文')
inds=0
for r in arg:
    inds=inds+1
    pap=PapersDictSelected[keys[r]]
#     if len(doiDict[pap['doi']]['name'])>1:
#         print(doiDict[pap['doi']]['name'])
#         continue;
    if numberingPapers:
        if pap['ryoiki']:
            p = document.add_paragraph(smark+str(inds)+'. '+pap['text1'])
        else:
            p = document.add_paragraph(str(inds)+'. '+pap['text1'])
    else:
        if pap['ryoiki']:
            p = document.add_paragraph(smark+pap['text1'])
        else:
            p = document.add_paragraph(pap['text1'])
    for nm in pap['authors']:
        if nm in nameList:
            if pap['Corresp']:
                p.add_run('*')
            if daihyobuntanList[nameList.index(nm)]=='D':
                p.add_run(nm).underline = WD_UNDERLINE.DOUBLE
            elif daihyobuntanList[nameList.index(nm)]=='B':
                p.add_run(nm).underline = True
            else:
                p.add_run(nm)
        else:
            p.add_run(nm)
        p.add_run(', ')
    p.add_run(pap['text2'])
    # if doiadd:
    #     p.add_run(' doi: '+pap['doi'])
    
TalksDictSelected={k:TalksDict[k] for k in range(len(TalksDict)) if (TalksDict[k]['date']>mindate) & (TalksDict[k]['date']<maxdate) }

keys=list(TalksDictSelected.keys())
datelist=[TalksDictSelected[r]['date'] for r in keys]
arg=np.argsort(datelist)[::-1]
document.add_paragraph('')
document.add_paragraph('学会発表・講演（招待あり）')
inds=0
for r in arg:
    inds=inds+1
    pap=TalksDictSelected[keys[r]]
    p = document.add_paragraph(str(inds)+'. ')
    nm=pap["presenter"]
    if daihyobuntanList[nameList.index(nm)]=='D':
        p.add_run(pap["printname"]).underline = WD_UNDERLINE.DOUBLE
    elif daihyobuntanList[nameList.index(nm)]=='B':
        p.add_run(pap["printname"]).underline = True
    else:
        p.add_run(pap["printname"])
    p.add_run(', \"'+pap["presentation_title"]+"\"")
    p.add_run(', '+pap["event"])
    p.add_run(', '+pap["date"]+'.')

booksDictSelected={k:booksDict[k] for k in range(len(booksDict)) if (booksDict[k]['date']>mindate) & (booksDict[k]['date']<maxdate) }

keys=list(booksDictSelected.keys())
datelist=[booksDictSelected[r]['date'] for r in keys]
arg=np.argsort(datelist)[::-1]
document.add_paragraph('')
document.add_paragraph('書籍')
inds=0
for r in arg:
    inds=inds+1
    pap=booksDictSelected[keys[r]]
    p = document.add_paragraph(str(inds)+'. ')
    nm=pap['authors']
    if daihyobuntanList[nameList.index(nm)]=='D':
        p.add_run(pap["printname"]).underline = WD_UNDERLINE.DOUBLE
    elif daihyobuntanList[nameList.index(nm)]=='B':
        p.add_run(pap["printname"]).underline = True
    else:
        p.add_run(pap["printname"])
    p.add_run(pap["book_owner_role"])
    p.add_run(pap["book_owner_range"])
    p.add_run(pap["book_title"])
    p.add_run(pap["publisher"])
    p.add_run(' '+pap["date"][:7]+'.')

miscDictSelected={k:miscDict[k] for k in range(len(miscDict)) if (miscDict[k]['date']>mindate) & (miscDict[k]['date']<maxdate) }

keys=list(miscDictSelected.keys())
datelist=[miscDictSelected[r]['date'] for r in keys]
arg=np.argsort(datelist)[::-1]
document.add_paragraph('')
document.add_paragraph('その他')
inds=0
for r in arg:
    inds=inds+1
    pap=miscDictSelected[keys[r]]
    p = document.add_paragraph(str(inds)+'. ')
    nm=pap['authors']
    if daihyobuntanList[nameList.index(nm)]=='D':
        p.add_run(pap["printname"]).underline = WD_UNDERLINE.DOUBLE
    elif daihyobuntanList[nameList.index(nm)]=='B':
        p.add_run(pap["printname"]).underline = True
    else:
        p.add_run(pap["printname"])
    p.add_run(','+pap["paper_title"])
    p.add_run(pap["publication_name"])
    p.add_run(' '+pap["date"]+'.')
document.save(file_name)

# %%
if ('google.colab' in str(get_ipython())):
    files.download(file_name)


